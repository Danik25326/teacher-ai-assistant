import os
from docx import Document
from typing import Optional, List

def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Витягує текст із .docx файлу."""
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Помилка читання {file_path}: {e}")
        return None

def extract_topics_from_docx(file_path: str) -> List[str]:
    """
    Найпростіше витягування тем (заголовків) із документа.
    Можна покращити під свої підручники.
    """
    topics = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and (text.isupper() or any(ch.isdigit() for ch in text[:5])):
                topics.append(text)
    except:
        pass
    return topics[:20]

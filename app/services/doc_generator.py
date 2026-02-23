from docx import Document
from docx.shared import Inches
from app.models.lesson import LessonPlan

def create_lesson_doc(lesson: LessonPlan, filename: str):
    """Створює документ Word із конспектом уроку."""
    doc = Document()
    doc.add_heading(f'Конспект уроку: {lesson.topic}', level=0)

    doc.add_heading('Конспект', level=1)
    doc.add_paragraph(lesson.summary)

    doc.add_heading('Завдання на уроці', level=1)
    for task in lesson.classwork:
        doc.add_paragraph(f'• {task}', style='List Bullet')

    doc.add_heading('Домашнє завдання', level=1)
    for task in lesson.homework:
        doc.add_paragraph(f'• {task}', style='List Bullet')

    doc.save(filename)
    return filename
